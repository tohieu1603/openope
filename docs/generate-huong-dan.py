"""Generate huong-dan-su-dung.docx — Hướng dẫn sử dụng Operis cho người dùng."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

SCRIPT_DIR = os.path.dirname(__file__)
IMG_DIR = os.path.join(SCRIPT_DIR, "screenshots")

doc = Document()

# ── Styles ──────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

BLUE = RGBColor(0x1B, 0x3A, 0x5C)       # heading color
DARK = RGBColor(0x2D, 0x2D, 0x2D)       # body text
CAPTION_GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x0D, 0x6E, 0xAA)     # links / highlights
SEPARATOR_COLOR = RGBColor(0xBB, 0xBB, 0xBB)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Arial"
    h.font.color.rgb = BLUE

doc.styles["Heading 1"].font.size = Pt(18)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 3"].font.size = Pt(12)

# List Bullet style
if "List Bullet" not in [s.name for s in doc.styles]:
    doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)

# Caption style
cap_style = doc.styles.add_style("ImageCaption", WD_STYLE_TYPE.PARAGRAPH)
cap_style.font.name = "Arial"
cap_style.font.size = Pt(10)
cap_style.font.italic = True
cap_style.font.color.rgb = CAPTION_GRAY
cap_style.paragraph_format.space_before = Pt(4)
cap_style.paragraph_format.space_after = Pt(12)
cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

img_counter = [0]


# ── Helpers ─────────────────────────────────────────────────

def title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = BLUE


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.color.rgb = DARK


def bold_para(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.color.rgb = DARK


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.color.rgb = DARK
        r2 = p.add_run(item)
        r2.font.color.rgb = DARK


def img(filename, caption):
    """Chèn ảnh thật từ screenshots/ và thêm chú thích bên dưới."""
    img_counter[0] += 1
    path = os.path.join(IMG_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
    else:
        # Placeholder nếu ảnh chưa có
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Chưa có ảnh: {filename}]")
        r.font.size = Pt(10)
        r.font.color.rgb = CAPTION_GRAY
    # Chú thích
    doc.add_paragraph(
        f"Hình {img_counter[0]}: {caption}", style="ImageCaption"
    )


def separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("―" * 40)
    r.font.color.rgb = SEPARATOR_COLOR


def qa(question, answer):
    p = doc.add_paragraph()
    r = p.add_run(f"Hỏi: {question}")
    r.bold = True
    r.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Trả lời: {answer}")
    r2.font.color.rgb = DARK


# ════════════════════════════════════════════════════════════
# NỘI DUNG
# ════════════════════════════════════════════════════════════

title("Hướng Dẫn Sử Dụng Operis")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Tài liệu dành cho người dùng. Không cần biết về lập trình.")
r.font.color.rgb = CAPTION_GRAY
r.font.size = Pt(11)

separator()

# ── MỤC LỤC ──
h1("Mục lục")
toc = [
    "Giới thiệu về Operis",
    "Truy cập trang web operis.vn",
    "Đăng ký tài khoản",
    "Đăng nhập tài khoản",
    "Nạp token (nạp tiền sử dụng)",
    "Cổng quản lý Operis Client",
    "Chat - Trò chuyện với trợ lý AI",
    "Thanh toán - Nạp tiền và quản lý số dư",
    "Phân tích - Xem thống kê sử dụng",
    "Nhật ký - Lịch sử hội thoại",
    "Workflow - Tự động hoá công việc",
    "Tài liệu - Hướng dẫn chi tiết",
    "Góp ý - Báo lỗi và đề xuất",
    "Cài đặt - Quản lý tài khoản",
    "Câu hỏi thường gặp",
]
for i, item in enumerate(toc, 1):
    p = doc.add_paragraph(f"{i}. {item}")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# PHẦN A: OPERIS.VN
# ════════════════════════════════════════════════════════════

h1("1. Giới thiệu về Operis")

para(
    "Operis là trợ lý AI cá nhân, được thiết kế để hỗ trợ bạn trong công việc hàng ngày "
    "như trả lời câu hỏi, soạn văn bản, tìm kiếm thông tin, dịch thuật và nhiều tác vụ khác. "
    "Bạn có thể trò chuyện với Operis qua trang web hoặc qua các ứng dụng nhắn tin "
    "như Zalo, Telegram, WhatsApp."
)

para(
    "Để bắt đầu sử dụng, bạn cần thực hiện các bước sau theo thứ tự:"
)

numbered([
    "Truy cập trang web operis.vn để tìm hiểu về sản phẩm.",
    "Đăng ký tài khoản.",
    "Đăng nhập.",
    "Nạp token (nạp tiền) để sử dụng dịch vụ.",
    "Truy cập Cổng quản lý để bắt đầu trò chuyện với trợ lý AI.",
])

separator()

# ── 2. TRANG WEB OPERIS.VN ──
h1("2. Truy cập trang web operis.vn")

para(
    "Mở trình duyệt web (Chrome, Safari, Firefox, ...) và truy cập địa chỉ:"
)
p = doc.add_paragraph()
r = p.add_run("https://operis.vn")
r.bold = True
r.font.color.rgb = ACCENT
r.font.size = Pt(12)

img("01-trang-chu.png", "Trang chủ operis.vn")

para(
    "Đây là trang chủ của Operis. Tại đây, bạn có thể xem thông tin giới thiệu về sản phẩm, "
    "các gói dịch vụ, và truy cập các trang quan trọng khác."
)

h2("2.1. Cửa hàng")

para(
    "Nhấn vào mục \"Cửa hàng\" trên thanh menu để xem danh sách các sản phẩm Operis."
)

img("02-cua-hang.png", "Trang cửa hàng - danh sách các gói sản phẩm Operis")

para(
    "Tại đây bạn có thể xem các gói sản phẩm với mức giá khác nhau. "
    "Nhấn vào từng sản phẩm để xem thông tin chi tiết."
)

img("02c-san-pham.png", "Trang chi tiết sản phẩm Operisbot Cá Nhân 1")

h2("2.2. Giới thiệu và Liên hệ")

para(
    "Bạn có thể nhấn vào \"Giới thiệu\" để tìm hiểu thêm về Operis, "
    "hoặc \"Liên hệ\" nếu cần hỗ trợ."
)

img("03-gioi-thieu.png", "Trang giới thiệu về Operis")

img("04-lien-he.png", "Trang liên hệ - thông tin hỗ trợ khách hàng")

separator()

# ── 3. ĐĂNG KÝ TÀI KHOẢN ──
h1("3. Đăng ký tài khoản")

para("Để sử dụng Operis, bạn cần có tài khoản. Các bước đăng ký như sau:")

numbered([
    "Truy cập trang operis.vn.",
    "Nhấn nút \"Đăng nhập\" ở góc trên bên phải.",
    "Tại trang đăng nhập, nhấn vào liên kết \"Đăng ký\" (hoặc \"Tạo tài khoản mới\").",
    "Điền đầy đủ thông tin: Họ và tên, Email, Mật khẩu (tối thiểu 8 ký tự), Xác nhận mật khẩu.",
    "Nhấn nút \"Đăng ký\".",
    "Hệ thống sẽ gửi email xác nhận. Vui lòng kiểm tra hộp thư và xác nhận tài khoản.",
])

img("05d-operis-register.png", "Trang đăng ký tài khoản tại operis.vn")

bold_para("Lưu ý: ", "Vui lòng nhớ kỹ email và mật khẩu đã đăng ký để đăng nhập sau này.")

separator()

# ── 4. ĐĂNG NHẬP ──
h1("4. Đăng nhập tài khoản")

para("Sau khi đã có tài khoản, bạn đăng nhập để sử dụng dịch vụ.")

numbered([
    "Truy cập operis.vn và nhấn nút \"Đăng nhập\" ở góc trên bên phải.",
    "Nhập Email đã đăng ký.",
    "Nhập Mật khẩu.",
    "Nhấn nút \"Đăng nhập\".",
])

img("05-operis-login.png", "Trang đăng nhập operis.vn - nhập Email và Mật khẩu")

img("05b-operis-login-filled.png", "Trang đăng nhập sau khi đã điền thông tin")

para(
    "Sau khi đăng nhập thành công, bạn sẽ được chuyển đến trang quản lý tài khoản."
)

img("05c-operis-after-login.png", "Trang tài khoản sau khi đăng nhập thành công")

para(
    "Tại trang tài khoản, bạn có thể xem và quản lý:"
)
bullet("Hồ sơ cá nhân")
bullet("Token và Gói dịch vụ")
bullet("Lịch sử sử dụng API")
bullet("Đơn hàng")
bullet("Lịch sử giao dịch")

separator()

# ── 5. NẠP TOKEN ──
h1("5. Nạp token (nạp tiền sử dụng)")

para(
    "Token là đơn vị sử dụng dịch vụ Operis. Mỗi khi bạn gửi tin nhắn cho trợ lý AI, "
    "hệ thống sẽ trừ một số token tương ứng. Để tiếp tục sử dụng, bạn cần nạp thêm token."
)

h2("5.1. Các gói nạp token")

para("Operis cung cấp nhiều gói nạp token phù hợp với nhu cầu khác nhau:")

h3("Gói Thuê bao")
para("Gói cơ bản với chi phí hợp lý, phù hợp cho người dùng cá nhân có nhu cầu sử dụng vừa phải.")

img("06a-nap-token-thue-bao.png", "Gói Thuê bao - gói nạp token cơ bản")

h3("Gói Tháng 1")
para("Gói tiêu chuẩn với lượng token lớn hơn, phù hợp cho người dùng thường xuyên.")

img("06b-nap-token-thang1.png", "Gói Tháng 1 - gói nạp token tiêu chuẩn")

h3("Gói VIP")
para("Gói cao cấp với lượng token lớn nhất, dành cho người dùng chuyên nghiệp hoặc doanh nghiệp nhỏ.")

img("06c-nap-token-vip.png", "Gói VIP - gói nạp token cao cấp")

h2("5.2. Cách nạp token")

numbered([
    "Đăng nhập tài khoản tại operis.vn.",
    "Vào phần \"Token & Gói dịch vụ\" trong trang tài khoản.",
    "Chọn gói nạp phù hợp với nhu cầu.",
    "Nhấn nút \"Mua\" hoặc \"Nạp token\".",
    "Hệ thống sẽ hiển thị mã QR thanh toán.",
    "Mở ứng dụng ngân hàng trên điện thoại, quét mã QR và xác nhận thanh toán.",
    "Sau khi thanh toán thành công, token sẽ được cộng vào tài khoản tự động.",
])

bold_para("Lưu ý: ", "Token đã nạp không có thời hạn sử dụng. Ngoài ra, hệ thống cũng cấp miễn phí một lượng token nhỏ định kỳ mỗi 5 giờ.")

separator()

# ════════════════════════════════════════════════════════════
# PHẦN B: CỔNG QUẢN LÝ (CLIENT-WEB)
# ════════════════════════════════════════════════════════════

h1("6. Cổng quản lý Operis Client")

para(
    "Cổng quản lý Operis Client là nơi bạn sử dụng trực tiếp trợ lý AI và quản lý tài khoản. "
    "Sau khi mua sản phẩm Operis, quản trị viên sẽ cung cấp cho bạn đường dẫn truy cập."
)

h2("6.1. Đăng nhập Cổng quản lý")

numbered([
    "Mở trình duyệt và truy cập đường dẫn do quản trị viên cung cấp.",
    "Nhập Email và Mật khẩu.",
    "Nhấn nút \"Đăng nhập\".",
])

img("10-client-login.png", "Trang đăng nhập Cổng quản lý Operis Client")

h2("6.2. Giao diện chính")

para(
    "Sau khi đăng nhập, bạn sẽ thấy giao diện chính gồm hai phần:"
)

bold_para("Thanh menu bên trái: ", "Chứa các mục chức năng. Nhấn vào từng mục để chuyển trang.")
bold_para("Vùng nội dung bên phải: ", "Hiển thị nội dung của mục bạn đang chọn.")

para("Các mục trong menu bao gồm:")
bullet("Chat - Trò chuyện trực tiếp với trợ lý AI")
bullet("Phân tích - Xem thống kê sử dụng")
bullet("Thanh toán - Nạp tiền, xem số dư")
bullet("Nhật ký - Xem lại lịch sử hội thoại")
bullet("Workflow - Đặt lịch tự động hoá công việc")
bullet("Tài liệu - Đọc tài liệu hướng dẫn chi tiết")
bullet("Góp ý - Báo lỗi hoặc gửi đề xuất")
bullet("Cài đặt - Cài đặt tài khoản, kết nối ứng dụng nhắn tin")

separator()

# ── 7. CHAT ──
h1("7. Chat - Trò chuyện với trợ lý AI")

para(
    "Đây là chức năng chính của Operis. Bạn gõ tin nhắn và trợ lý AI sẽ trả lời."
)

img("12-chat.png", "Trang Chat - danh sách hội thoại bên trái, khung chat bên phải")

h2("7.1. Cách sử dụng")

numbered([
    "Nhấn vào \"Chat\" trên menu.",
    "Gõ tin nhắn vào ô ở cuối trang.",
    "Nhấn Enter hoặc nút Gửi.",
    "Đợi vài giây, trợ lý sẽ trả lời.",
])

h2("7.2. Bạn có thể hỏi gì?")

bullet("Hỏi đáp: \"Thời tiết hôm nay thế nào?\", \"GDP Việt Nam năm 2024?\"")
bullet("Nhờ viết: \"Viết email xin nghỉ phép ngày mai\", \"Viết bài giới thiệu sản phẩm\"")
bullet("Dịch thuật: \"Dịch đoạn này sang tiếng Anh: ...\"")
bullet("Tính toán: \"100 USD bằng bao nhiêu VND?\"")
bullet("Tìm kiếm: \"Tìm 5 nhà hàng ngon ở Quận 1\"")
bullet("Tóm tắt: \"Tóm tắt bài viết này trong 3 dòng\"")

h2("7.3. Quản lý cuộc hội thoại")

bullet("Tạo hội thoại mới: Nhấn nút + để bắt đầu chủ đề mới.")
bullet("Chuyển đổi: Nhấn vào cuộc hội thoại cũ trong danh sách bên trái để xem lại.")
bullet("Xoá: Nhấn biểu tượng thùng rác để xoá cuộc hội thoại không cần nữa.")

h2("7.4. Mẹo sử dụng")

bullet("Nói càng cụ thể càng tốt. Ví dụ: \"Tìm giá iPhone 16 Pro Max trên Shopee\" tốt hơn \"tìm điện thoại\".")
bullet("Trợ lý nhớ nội dung cuộc hội thoại, bạn có thể hỏi tiếp mà không cần nhắc lại.")
bullet("Nếu kết quả chưa đúng ý, hãy nói lại cụ thể hơn.")

separator()

# ── 8. THANH TOÁN ──
h1("8. Thanh toán - Nạp tiền và quản lý số dư")

para("Trang này giúp bạn xem số dư, nạp tiền và theo dõi lịch sử thanh toán.")

img("14-billing.png", "Trang Thanh toán - hiển thị số dư token và khu vực nạp tiền")

h2("8.1. Xem số dư")

para("Ở đầu trang, bạn thấy 2 loại số dư:")
bullet("Số dư chính (token trả phí): Số tiền bạn đã nạp, trừ dần khi sử dụng.")
bullet("Số dư miễn phí: Token miễn phí được cấp định kỳ mỗi 5 giờ.")

h2("8.2. Nạp tiền")

numbered([
    "Chọn gói nạp phù hợp (hoặc nhập số tiền tuỳ chỉnh).",
    "Nhấn nút \"Mua Tokens\".",
    "Màn hình hiện mã QR để thanh toán.",
    "Mở ứng dụng ngân hàng trên điện thoại.",
    "Quét mã QR và xác nhận thanh toán.",
    "Đợi vài giây, hệ thống sẽ tự xác nhận và cộng token vào tài khoản.",
])

bold_para("Lưu ý: ", "Nếu hệ thống chưa tự nhận, vui lòng nhấn nút \"Kiểm tra giao dịch\" để kiểm tra lại.")

h2("8.3. Lịch sử giao dịch")

para("Cuộn xuống dưới để xem danh sách các giao dịch đã thực hiện:")
bullet("Chờ thanh toán: Giao dịch đang đợi bạn thanh toán.")
bullet("Hoàn thành: Đã thanh toán thành công.")
bullet("Đã huỷ: Giao dịch bị huỷ.")
bullet("Hết hạn: Quá thời gian thanh toán.")

separator()

# ── 9. PHÂN TÍCH ──
h1("9. Phân tích - Xem thống kê sử dụng")

para("Trang này cho bạn biết bạn đã dùng trợ lý bao nhiêu.")

img("13-analytics.png", "Trang Phân tích - biểu đồ sử dụng và các con số thống kê")

h2("9.1. Thông tin hiển thị")

bullet("Số token đã dùng: Tổng lượng token tiêu thụ trong khoảng thời gian.")
bullet("Số lượt sử dụng: Bao nhiêu lần bạn đã gửi tin nhắn.")
bullet("Biểu đồ theo ngày: Xem mức độ sử dụng từng ngày.")

h2("9.2. Chọn khoảng thời gian")

para("Nhấn vào các nút ở trên để chọn:")
bullet("1 ngày: Chỉ xem hôm nay.")
bullet("7 ngày: Tuần vừa qua.")
bullet("30 ngày: Tháng vừa qua.")
bullet("90 ngày: 3 tháng gần nhất.")
bullet("Tuỳ chỉnh: Chọn ngày bắt đầu và kết thúc theo ý bạn.")

separator()

# ── 10. NHẬT KÝ ──
h1("10. Nhật ký - Lịch sử hội thoại")

para("Xem lại tất cả các cuộc hội thoại trước đó.")

img("15-logs.png", "Trang Nhật ký - danh sách cuộc hội thoại với ngày giờ")

h2("10.1. Cách sử dụng")

numbered([
    "Nhấn \"Nhật ký\" trên menu.",
    "Danh sách hiện ra với các cuộc hội thoại gần nhất.",
    "Nhấn vào cuộc hội thoại để xem chi tiết.",
])

h2("10.2. Tìm kiếm")

bullet("Gõ từ khoá vào ô tìm kiếm ở trên cùng để tìm cuộc hội thoại cụ thể.")
bullet("Ví dụ: gõ \"báo cáo\" để tìm các cuộc hội thoại liên quan đến báo cáo.")

separator()

# ── 11. WORKFLOW ──
h1("11. Workflow - Tự động hoá công việc")

para(
    "Workflow giúp bạn đặt lịch cho trợ lý tự động làm việc vào thời gian cố định."
)

img("16-workflow.png", "Trang Workflow - tạo và quản lý công việc tự động")

h2("11.1. Workflow là gì?")

para(
    "Hãy tưởng tượng bạn đặt hẹn giờ: \"Mỗi sáng 8h, tìm thời tiết và gửi cho tôi\". "
    "Trợ lý sẽ tự động làm điều đó mà bạn không cần nhắc."
)

h2("11.2. Tạo Workflow mới")

numbered([
    "Điền tên cho workflow (ví dụ: \"Thời tiết buổi sáng\").",
    "Chọn lịch chạy: mỗi ngày, mỗi tuần, hoặc mỗi bao nhiêu phút/giờ.",
    "Nhập nội dung yêu cầu (ví dụ: \"Tìm thời tiết TP.HCM hôm nay\").",
    "Nhấn \"Lưu\".",
])

h2("11.3. Quản lý Workflow")

bullet("Bật/Tắt: Nhấn công tắc để tạm dừng hoặc kích hoạt lại.")
bullet("Chạy ngay: Nhấn nút chạy để thực hiện ngay mà không đợi đến lịch.")
bullet("Xoá: Nhấn nút xoá để loại bỏ workflow không cần.")

h2("11.4. Ví dụ thực tế")

bullet("\"Mỗi sáng 7h tìm thời tiết và gửi cho tôi\"")
bullet("\"Mỗi 10 phút kiểm tra trang web của tôi có hoạt động không\"")
bullet("\"Mỗi tối 8h tổng hợp tin tức công nghệ\"")
bullet("\"Mỗi thứ 2 lúc 9h nhắc tôi gửi báo cáo tuần\"")

separator()

# ── 12. TÀI LIỆU ──
h1("12. Tài liệu - Hướng dẫn chi tiết")

para("Trang này chứa các bài hướng dẫn chi tiết về cách sử dụng từng tính năng.")

img("17-docs.png", "Trang Tài liệu - các danh mục hướng dẫn")

h2("12.1. Cách sử dụng")

numbered([
    "Nhấn \"Tài liệu\" trên menu.",
    "Chọn danh mục bạn quan tâm (ví dụ: \"Khả Năng Cơ Bản\").",
    "Nhấn vào bài viết cụ thể để đọc.",
])

h2("12.2. Các danh mục")

bullet("Giới thiệu: Bắt đầu nhanh, tổng quan hệ thống.")
bullet("Khả năng cơ bản: Chat, kiểm tra máy tính, quản lý file, tìm kiếm web.")
bullet("Trình duyệt web: Điều khiển trình duyệt, cài Chrome Extension.")
bullet("Tự động hoá: Hẹn giờ, nhắc nhở, kiểm tra định kỳ.")

separator()

# ── 13. GÓP Ý ──
h1("13. Góp ý - Báo lỗi và đề xuất")

para("Khi gặp lỗi hoặc có ý tưởng cải thiện, bạn gửi phản hồi tại đây.")

img("18-gop-y.png", "Trang Góp ý - biểu mẫu gửi phản hồi")

h2("13.1. Gửi góp ý")

numbered([
    "Chọn loại phản hồi:",
])
bullet("\"Báo lỗi\": Khi bạn gặp lỗi, trợ lý không hoạt động đúng.")
bullet("\"Góp ý\": Ý kiến cải thiện chất lượng dịch vụ.")
bullet("\"Ý tưởng\": Đề xuất tính năng mới.")

para("Tiếp theo:")
numbered([
    "Nhập tiêu đề ngắn gọn.",
    "Mô tả chi tiết vấn đề hoặc đề xuất.",
    "Nhấn \"Gửi\".",
])

h2("13.2. Theo dõi trạng thái")

para("Sau khi gửi, bạn có thể xem trạng thái phản hồi:")
bullet("Mở: Đã nhận, đang chờ xem xét.")
bullet("Đang xử lý: Đội ngũ đang giải quyết.")
bullet("Đã giải quyết: Vấn đề đã được khắc phục.")
bullet("Đóng: Phản hồi đã được đóng.")

separator()

# ── 14. CÀI ĐẶT ──
h1("14. Cài đặt - Quản lý tài khoản")

para("Quản lý thông tin cá nhân, bảo mật và kết nối ứng dụng nhắn tin.")

img("19-settings.png", "Trang Cài đặt - thông tin cá nhân, kết nối ứng dụng, bảo mật")

h2("14.1. Thông tin cá nhân")

bullet("Xem và sửa tên hiển thị: Nhấn nút sửa bên cạnh tên, gõ tên mới, nhấn Lưu.")

h2("14.2. Kết nối ứng dụng nhắn tin")

para(
    "Bạn có thể kết nối trợ lý với các ứng dụng nhắn tin quen thuộc để trò chuyện "
    "trực tiếp trên đó, thay vì phải mở trang web."
)

para("Các ứng dụng được hỗ trợ:")
bullet("WhatsApp")
bullet("Telegram")
bullet("Zalo")

h3("Cách kết nối (ví dụ với Zalo)")

numbered([
    "Nhấn nút \"Kết nối\" bên cạnh Zalo.",
    "Màn hình hiện mã QR.",
    "Mở Zalo trên điện thoại.",
    "Quét mã QR.",
    "Xác nhận trên điện thoại.",
    "Khi kết nối thành công, trạng thái chuyển sang \"Đã kết nối\".",
])

para(
    "Sau khi kết nối, bạn mở Zalo (hoặc ứng dụng đã kết nối) và nhắn tin trực tiếp "
    "- trợ lý sẽ trả lời ngay trên đó."
)

bold_para("Ngắt kết nối: ", "Nhấn nút \"Ngắt kết nối\" nếu bạn không muốn dùng kênh đó nữa.")

h2("14.3. Bảo mật - Đổi mật khẩu")

numbered([
    "Nhấn vào \"Đổi mật khẩu\".",
    "Nhập Mật khẩu hiện tại.",
    "Nhập Mật khẩu mới.",
    "Nhấn \"Xác nhận\".",
])

separator()

# ── 15. CÂU HỎI THƯỜNG GẶP ──
h1("15. Câu hỏi thường gặp")

qa(
    "Trợ lý có hiểu tiếng Việt không?",
    "Có. Bạn cứ nhắn tin bằng tiếng Việt bình thường, như nhắn tin Zalo vậy."
)

qa(
    "Tôi dùng hết token thì sao?",
    "Bạn sẽ không gửi được tin nhắn mới. Hãy vào trang Thanh toán để nạp thêm. "
    "Ngoài ra, hệ thống cấp miễn phí một lượng token nhỏ mỗi 5 giờ."
)

qa(
    "Dữ liệu của tôi có an toàn không?",
    "Có. Hệ thống chạy trên máy chủ riêng, dữ liệu không bị chia sẻ ra ngoài."
)

qa(
    "Trợ lý có sai không?",
    "Có thể. Giống như hỏi một người, đôi khi câu trả lời chưa chính xác 100%. "
    "Nên kiểm tra lại thông tin quan trọng."
)

qa(
    "Tôi có thể dùng trên điện thoại không?",
    "Có. Mở trình duyệt trên điện thoại, vào đường dẫn do quản trị viên cung cấp. "
    "Giao diện tự động điều chỉnh cho màn hình nhỏ."
)

qa(
    "Kết nối WhatsApp/Telegram/Zalo có mất phí không?",
    "Không. Kết nối miễn phí. Bạn chỉ tốn token khi gửi tin nhắn (dù qua kênh nào)."
)

qa(
    "Workflow là gì, có cần biết lập trình không?",
    "Không cần. Workflow giống như đặt hẹn giờ cho trợ lý. Bạn chỉ cần gõ nội dung "
    "yêu cầu bằng tiếng Việt bình thường."
)

qa(
    "Liên hệ hỗ trợ ở đâu?",
    "Vào trang \"Góp ý\" để gửi báo lỗi, hoặc liên hệ qua trang \"Liên hệ\" tại operis.vn."
)

# Footer
separator()
p = doc.add_paragraph("Tài liệu cập nhật: Tháng 2/2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.color.rgb = CAPTION_GRAY

p2 = doc.add_paragraph("Operis - Trợ lý AI cá nhân của bạn")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p2.runs:
    run.font.color.rgb = BLUE
    run.bold = True

# ── Save ──
out = os.path.join(SCRIPT_DIR, "operis.docx")
doc.save(out)
print(f"OK → {out}")
